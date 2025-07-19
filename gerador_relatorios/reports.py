from __future__ import annotations

import os
from typing import Dict, Any
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from math import ceil
from docx.enum.table import WD_TABLE_ALIGNMENT

from .utils import resource_path

SIGNATURES: dict[str, list[tuple[str, dict[str, bool]]]] = {
    "PSICOTERAPIA": [
        ("assinado eletronicamente", {"italic": True}),
        ("SIMONE ULER LAVORATO", {"bold": True}),
        ("Psicóloga", {}),
        ("CRP 01/20631", {})
    ],
    "ABA": [
        ("assinado eletronicamente", {"italic": True}),
        ("SIMONE ULER LAVORATO", {"bold": True}),
        ("Psicóloga", {}),
        ("CRP 01/20631", {})
    ],
    "TERAPIA OCUPACIONAL": [
        ("assinado eletronicamente", {"italic": True}),
        ("WÊNIA CARVALHO TORRES", {"bold": True}),
        ("Terapeuta Ocupacional", {}),
        ("CREFITO 11-17626", {})
    ],
    "FONOAUDIOLOGIA": [
        ("assinado eletronicamente", {"italic": True}),
        ("GABRIEL BORGES BATISTA", {"bold": True}),
        ("Fonoaudiólogo", {}),
        ("CRFa 5-13809", {})
    ],
    "PSICOMOTRICIDADE": [
        ("assinado eletronicamente", {"italic": True}),
        ("SIMONE ULER LAVORATO", {"bold": True}),
        ("Psicomotricista", {}),
        ("CRP 01/20631", {})
    ],
    "PSICOPEDAGOGIA": [
        ("assinado eletronicamente", {"italic": True}),
        ("SIMONE ULER LAVORATO", {"bold": True}),
        ("Psicopedagoga", {}),
        ("CRP 01/20631", {})
    ],
    "NUTRIÇÃO": [
        ("assinado eletronicamente", {"italic": True}),
        ("RAIANE ALVES ROCHA", {"bold": True}),
        ("Nutricionista", {}),
        ("CRN 17753", {})
    ]
}

# Métodos para moldar os templates

def add_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for border_name in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "000000")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def create_header_table(doc: Document, patient_data: Dict[str, Any], convenio: str = "FUSEX"):
    table = doc.add_table(rows=6, cols=1)
    add_table_borders(table)

    especialidades_unicas = list(set(patient_data["especialidades"]))
    especialidades_str = ", ".join(especialidades_unicas)

    header_data = [
        f"Nome: {patient_data['info']['nome']}",
        f"Data de Nascimento: {patient_data['info']['data_nascimento']}",
        f"Responsável: {patient_data['info']['responsavel']}",
        f"Convênio: {convenio}",
        f"Especialidade: {especialidades_str}",
        f"Mês de referência: {patient_data['info']['mes_referencia']}",
    ]

    for i, data in enumerate(header_data):
        cell = table.cell(i, 0)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if ":" in data:
            parts = data.split(":", 1)
            title_run = paragraph.add_run(parts[0] + ": ")
            title_run.bold = True
            paragraph.add_run(parts[1])
        else:
            run = paragraph.add_run(data)
            run.bold = True

    doc.add_paragraph()
    return table


def add_section_title(doc: Document, title: str, space_before: Pt = Pt(24), space_after: Pt = Pt(12)):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = space_before
    fmt.space_after = space_after
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(title)
    run.bold = True
    return p


def add_section_text(doc: Document, text: str, space_before: Pt = Pt(12), space_after: Pt = Pt(12)):
    p = doc.add_paragraph(text)
    fmt = p.paragraph_format
    fmt.space_before = space_before
    fmt.space_after = space_after
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_specialty_section(
    doc: Document,
    specialty_name: str,
    evolution_text: str,
    programming_text: str | None = None,
    is_evolution: bool = True,
):
    if is_evolution:
        add_section_title(doc, specialty_name, Pt(18), Pt(8))
    else:
        add_section_title(doc, specialty_name, Pt(16), Pt(8))

    add_section_text(doc, evolution_text)

    if programming_text:
        add_section_text(doc, programming_text)

def add_fixed_signature_section(
        doc: Document,
        text: str = "Brasília, data da assinatura digital.",
        font_size: int = 12,
        alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.RIGHT,
) -> None: 
    """Insere uma linha fixa com cidade e texto de data nates da assinatura digital."""
    p = doc.add_paragraph(text)
    p.alignment = alignment
    run = p.runs[0]
    run.font.size = Pt(font_size)

# Depreciado, use `build_signature_block`.
def add_signature_section(doc: Document, signature_type: str = "tipico"):
    doc.add_paragraph()
    p = doc.add_paragraph("_" * 50)
    fmt = p.paragraph_format
    fmt.space_before = Pt(24)
    fmt.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Responsável técnica(o)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def build_signature_block_grid(
    doc: Document,
    specialties: list[str],
    cols: int = 2,
    font_size: int = 12
) -> None:
    """
    Insere:
      1) uma linha em branco
      2) tabela com `cols` colunas e linhas suficientes para todas as especialidades,
         preenchendo cada célula com o bloco de assinatura correspondente.
    """
    # 1) espaço antes
    doc.add_paragraph()

    # 3) monta lista de blocos só para as especialidades que batem
    upper_sps = [s.upper() for s in specialties]
    blocks = [
        SIGNATURES[esp]
        for esp in SIGNATURES
        if esp in upper_sps
    ]

    # 4) se nada bateu, usa um fallback único
    if not blocks:
        blocks = [[
            ("assinado eletronicamente", {"italic": True}),
            ("SIMONE ULER LAVORATO",     {"bold": True}),
            ("Psicóloga",                {}),
            ("CRP 01/20631",             {}),
        ]]

    # 5) calcula quantas linhas são necessárias
    rows = ceil(len(blocks) / cols)

    # 6) cria tabela  rows×cols
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 7) percorre cada bloco e preenche a célula correspondente
    for idx, sig_block in enumerate(blocks):
        row = idx // cols
        col = idx % cols
        cell = table.cell(row, col)

        # injeta cada linha de assinatura dentro da célula
        for text, style in sig_block:
            p = cell.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(font_size)
            if style.get("bold"):
                run.bold = True
            if style.get("italic"):
                run.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_unordered_list_with_styling(
        doc: Document,
        items: list[list[tuple[str, dict]]],
        indent_level: int = 0,
        space_before: Pt = Pt(6),
        space_after: Pt = Pt(6)
    ):
        for item_parts in items:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Pt(indent_level * 12)
            p.paragraph_format.space_before = space_before
            p.paragraph_format.space_after = space_after

            for text, style in item_parts:
                run = p.add_run(text)
                if style.get("bold"):
                    run.bold = True
                if style.get("italic"):
                    run.italic = True
                if style.get("underline"):
                    run.underline = True

# Geradores de relatórios templates

def generate_pne_report(patient_data: Dict[str, Any], output_dir: str) -> None:
    try:
        papel_timbrado = resource_path("papel timbrado.docx")
        doc = Document(papel_timbrado)
    except Exception:
        doc = Document()
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_p.add_run("CLÍNICA MÉDICA - PNE")
        run.bold = True
        run.font.size = Pt(16)
        doc.add_paragraph()

    create_header_table(doc, patient_data, "Fusex PNE")

    add_section_title(doc, "Hipótese Diagnóstica")
    add_section_text(
        doc,
        "Transtorno do Espectro Autista, conforme critérios do DSM-5. "
        "Apresentando prejuízos significativos na comunicação social e "
        "comportamentos restritos e repetitivos, necessitando de apoio "
        "substancial em múltiplas áreas do desenvolvimento. A intervenção "
        "nutricional é essencial para abordar questões alimentares específicas.",
    )

    especialidades_encontradas = [esp.upper() for esp in patient_data["especialidades"]]

    if any("NUTRI" in esp or "NUTRIÇÃO" in esp for esp in especialidades_encontradas):
        add_section_title(doc, "Evolução", Pt(16), Pt(8))
        add_section_text(
            doc,
            "Desde o início do acompanhamento em terapia alimentar, observa-se progresso gradual na aceitação de novos alimentos, bem como aumento da tolerância a diferentes texturas, cores e temperaturas. O paciente tem apresentado maior disponibilidade para explorar o ambiente alimentar de forma positiva, com redução de comportamentos de recusa extrema ou esquiva. Em contextos estruturados e com suporte terapêutico, verifica-se participação mais ativa durante exposições alimentares e maior engajamento nas atividades propostas. Em casos infantis, o uso de estratégias lúdicas, modelagem e reforçamento positivo tem sido eficaz na promoção de avanços. Em pacientes adolescentes e adultos, nota-se maior consciência sobre suas dificuldades e disposição para experimentar novas abordagens comportamentais e cognitivas ligadas à alimentação.",
        )
    
    if any("FISIO" in esp or "FISIOTERAPIA" in esp for esp in especialidades_encontradas):
        add_section_title(doc, "Evolução Terapêutica", Pt(16), Pt(8))
        add_section_text(
            doc,
            "Desde o início do acompanhamento fisioterapêutico, observa-se progressos graduais na adaptação a estímulos motores e na aceitação de novas abordagens corporais. O paciente demonstra:"
        )
        add_unordered_list_with_styling(
            [
                {"Melhora na coordenação motora global", {"bold": True}},
                {"e na participação em atividades estruturadas."},
            ],
            [
                {"Aumento da tolerância sensorial,", {"bold": True}},
                {"com menor resistência a toques e manipulações terapêuticas."},
            ],
            [
                {"Maior engajamento em exercícios posturais e proprioceptivos,", {"bold": True}},
                {"favorecendo a consciência corporal e o equilíbrio."}
            ],
            [
                {"Diminuição de comportamentos de esquiva e recusa extrema,", {"bold": True}},
                {"tornando-se mais receptivo ao contato físico e às estratégias de intervenção"}
            ]
        )
        add_section_text(
            doc,
            "Nos atendimentos infantis, a fisioterapia tem sido realizada por meio de estratégias lúdicas, modelagem motora e reforço positivo. Em pacientes adolescentes e adultos, nota-se um avanço na percepção das próprias dificuldades e maior disposição para experimentar técnicas adaptativas."
        )

    if any("ABA" in esp or "TERAPIA ABA" in esp for esp in especialidades_encontradas):
        add_specialty_section(
            doc,
            "Terapia ABA",
            "A paciente tem apresentado evolução gradual no manejo de "
            "comportamentos desafiadores e na aquisição de habilidades adaptativas. "
            "Destaca-se o progresso na capacidade de seguir instruções, maior "
            "participação em atividades estruturadas e aumento da comunicação "
            "funcional, com boa aceitação aos programas propostos e resposta "
            "positiva ao reforço positivo.",
            is_evolution=True,
        )

    if any("PSICOTERAPIA" in esp for esp in especialidades_encontradas):
        add_specialty_section(
            doc,
            "Psicoterapia",
            "Observa-se maior abertura da paciente ao vínculo terapêutico, "
            "com avanços na expressão emocional, identificação de sentimentos "
            "e melhora na tolerância a frustrações. Há desenvolvimento de "
            "estratégias internas de enfrentamento e maior consciência sobre "
            "suas próprias emoções e comportamentos, adequadas à sua faixa etária.",
            is_evolution=True,
        )

    if any("TERAPIA OCUPACIONAL" in esp or "OCUPACIONAL" in esp for esp in especialidades_encontradas):
        add_specialty_section(
            doc,
            "Terapia Ocupacional",
            "Houve progresso no desempenho ocupacional, especialmente nas "
            "áreas de autorregulação, coordenação motora e autonomia nas "
            "atividades diárias. A paciente apresenta melhor organização "
            "sensorial e maior engajamento em tarefas funcionais, tanto em "
            "contextos lúdicos quanto nas rotinas cotidianas.",
            is_evolution=True,
        )

    if any("FONOAUDIOLOGIA" in esp or "FONO" in esp for esp in especialidades_encontradas):
        add_specialty_section(
            doc,
            "Fonoaudiologia",
            "Verifica-se avanço significativo nas habilidades comunicativas, "
            "seja por meio da fala, linguagem alternativa ou recursos expressivos "
            "e receptivos. A paciente demonstra melhor compreensão de comandos, "
            "maior intenção comunicativa e expansão do vocabulário funcional, "
            "além de avanços na articulação e fluência, conforme a necessidade individual.",
            is_evolution=True,
        )

    if any("PSICOMOTRICIDADE" in esp or "PSICOMOTOR" in esp for esp in especialidades_encontradas):
        add_specialty_section(
            doc,
            "Psicomotricidade",
            "A evolução psicomotora inclui melhora na coordenação global e fina, "
            "organização espacial e equilíbrio. A paciente demonstra maior "
            "consciência corporal e controle motor, com progressos que refletem "
            "positivamente no comportamento, na atenção e na interação social "
            "durante as atividades terapêuticas.",
            is_evolution=True,
        )

    if any("PSICOPEDAGOGIA" in esp or "PEDAGOG" in esp for esp in especialidades_encontradas):
        add_specialty_section(
            doc,
            "Psicopedagogia",
            "A paciente apresentou melhora na atenção, concentração e interesse "
            "por atividades que envolvem linguagem, raciocínio lógico e habilidades "
            "acadêmicas. Observa-se avanço na memória de trabalho, organização do "
            "pensamento e capacidade de seguir sequências, contribuindo para o "
            "desempenho escolar ou acadêmico, conforme a faixa etária.",
            is_evolution=True,
        )

    add_section_title(doc, "Programação Terapêutica Atual", Pt(30), Pt(12))

    if any("FISIO" in esp or "FISIOTERAPIA" in esp for esp in especialidades_encontradas):
        add_section_text(
            doc,
            "O plano de intervenção fisioterapêutico visa:"
        )
        add_unordered_list_with_styling(
            [
                {"Ampliar o repertório motor e funcional,", {"bold": True}},
                {"promovendo maior independência nas atividades diárias."},
            ],
            [
                {"Reduzir a ansiedade associada a estímulos físicos,", {"bold": True}},
                {"facilitando a interação com o ambiente e com e objetos táteis."},
            ]
        )
        

    if any("ABA" in esp or "TERAPIA ABA" in esp for esp in especialidades_encontradas):
        add_section_title(doc, "Terapia ABA", Pt(16), Pt(8))
        add_section_text(
            doc,
            "A programação atual visa fortalecer comportamentos funcionais, "
            "ampliar a comunicação, promover autonomia nas rotinas e reduzir "
            "comportamentos de oposição, fuga ou auto estimulação. São utilizados "
            "programas personalizados de ensino por tentativas discretas, ensino "
            "naturalístico e treino de habilidades sociais.",
        )

    if any("PSICOTERAPIA" in esp for esp in especialidades_encontradas):
        add_section_title(doc, "Psicoterapia", Pt(16), Pt(8))
        add_section_text(
            doc,
            "Os objetivos terapêuticos incluem promover o autoconhecimento, "
            "a regulação emocional, o desenvolvimento da autoestima e o enfrentamento "
            "saudável de desafios, utilizando abordagens adequadas à idade "
            "(brincadeiras simbólicas, recursos visuais, técnicas cognitivas, entre outras).",
        )

    if any("TERAPIA OCUPACIONAL" in esp or "OCUPACIONAL" in esp for esp in especialidades_encontradas):
        add_section_title(doc, "Terapia Ocupacional", Pt(16), Pt(8))
        add_section_text(
            doc,
            "As intervenções atuais priorizam o desenvolvimento da independência "
            "em atividades de vida diária (AVDs), o planejamento motor e a integração "
            "sensorial. São propostas atividades lúdicas e funcionais, com adaptações "
            "conforme a faixa etária, para favorecer o desempenho ocupacional global.",
        )

    if any("FONOAUDIOLOGIA" in esp or "FONO" in esp for esp in especialidades_encontradas):
        add_section_title(doc, "Fonoaudiologia", Pt(16), Pt(8))
        add_section_text(
            doc,
            "O foco terapêutico envolve o aperfeiçoamento da linguagem oral e/ou "
            "alternativa, melhora na compreensão e expressão verbal, bem como o "
            "desenvolvimento das habilidades fonológicas e comunicativas. A intervenção "
            "considera o nível atual de linguagem e o contexto escolar, familiar e social da paciente.",
        )

    if any("PSICOMOTRICIDADE" in esp or "PSICOMOTOR" in esp for esp in especialidades_encontradas):
        add_section_title(doc, "Psicomotricidade", Pt(16), Pt(8))
        add_section_text(
            doc,
            "O trabalho psicomotor tem como meta promover o domínio do corpo no espaço, "
            "controle postural, lateralidade e coordenação em diferentes níveis. As sessões "
            "envolvem jogos, circuitos e desafios motores com objetivos específicos para "
            "aprimorar a integração sensório-motora.",
        )

    if any("PSICOPEDAGOGIA" in esp or "PEDAGOG" in esp for esp in especialidades_encontradas):
        add_section_title(doc, "Psicopedagogia", Pt(16), Pt(8))
        add_section_text(
            doc,
            "A atuação psicopedagógica busca estimular habilidades cognitivas e acadêmicas, "
            "com estratégias personalizadas para desenvolver leitura, escrita, lógica "
            "matemática e resolução de problemas. O plano também inclui o fortalecimento "
            "da autoestima escolar e o apoio no planejamento e organização do tempo.",
        )
    
    if any("NUTRI" in esp or "NUTRIÇÃO" in esp for esp in especialidades_encontradas):
        add_section_text(
            doc,
            "A intervenção segue com o objetivo de ampliar o repertório alimentar, reduzir a ansiedade associada à alimentação e favorecer a construção de uma relação mais funcional e saudável com os alimentos. As estratégias utilizadas incluem: exposição gradual a novos alimentos, dessensibilização sistemática, treino de habilidades de enfrentamento, uso de reforçadores, orientação nutricional (em parceria com outros profissionais) e envolvimento familiar. O planejamento terapêutico é individualizado, levando em consideração as preferências, o histórico alimentar e os fatores sensoriais e emocionais que impactam a alimentação do paciente. O vínculo terapêutico tem sido essencial para a manutenção do engajamento e para a superação de resistências naturais do processo.",
        )

    add_section_title(doc, "Considerações Finais", Pt(30), Pt(12))
    if not(any("NUTRI" in esp or "NUTRIÇÃO" in esp for esp in especialidades_encontradas)):
        add_section_text(
            doc,
            "A paciente segue em acompanhamento com evolução positiva. O trabalho "
            "conjunto entre as especialidades tem favorecido ganhos significativos e "
            "generalização das habilidades desenvolvidas para diferentes ambientes. "
            "Recomendamos a continuidade do atendimento terapêutico e o envolvimento "
            "da família e/ou escola no processo.",
        )
        add_section_text(
            doc,
            "Nos colocamos à disposição para esclarecimentos sobre o processo "
            "terapêutico, bem como para oferecer orientações e suporte sempre que "
            "necessário, respeitando os limites éticos da atuação clínica.",
        )
    else:
        add_section_text(
            doc,
            "Recomenda-se a continuidade da terapia alimentar, com envolvimento ativo da família e, quando possível, articulação com a equipe multidisciplinar (psicólogos, nutricionistas, fonoaudiólogos, terapeutas ocupacionais, entre outros). A consistência das intervenções em diferentes ambientes (casa, escola, clínica) é fundamental para a generalização dos progressos alcançados. O processo terapêutico tem se mostrado eficaz na promoção de avanços alimentares, fortalecimento da autonomia do paciente e melhoria na qualidade de vida.\n\nNos colocamos à disposição para esclarecimentos sobre o processo terapêutico, bem como para oferecer orientações e suporte sempre que necessário, respeitando os limites éticos da atuação clínica.",
        )

    add_fixed_signature_section(doc, "Brasília, data da assinatura digital.")

    # add_signature_section(doc, "pne")
    build_signature_block_grid(doc, patient_data["especialidades"], cols=2)

    filename = f"Relatório_PNE_{patient_data['info']['nome'].replace(' ', '_')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)


def generate_tipico_report(patient_data: Dict[str, Any], output_dir: str) -> None:
    if any("NUTRI" in esp or "NUTRIÇÃO" in esp for esp in patient_data["especialidades"]):
        raise ValueError("Fusex Típico não contempla NUTRIÇÃO.")
    
    if any("FISIO" in esp or "FISIOTERAPIA" in esp for esp in patient_data["especialidades"]):
        raise ValueError("Fusex Típico não contempla FISIOTERAPIA.")
    
    try:
        papel_timbrado = resource_path("papel timbrado.docx")
        doc = Document(papel_timbrado)
    except Exception:
        doc = Document()
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_p.add_run("CLÍNICA MÉDICA")
        run.bold = True
        run.font.size = Pt(16)
        doc.add_paragraph()

    create_header_table(doc, patient_data, "FUSEX")

    add_section_title(doc, "Hipótese Diagnóstica")
    add_section_text(
        doc,
        "O paciente apresenta características compatíveis com o desenvolvimento "
        "típico, sem indicação, até o momento, de transtornos diagnosticáveis "
        "conforme os manuais classificatórios vigentes (CID-11/DSM-5). As demandas "
        "observadas referem-se a dificuldades específicas no enfrentamento de "
        "situações do cotidiano, que podem envolver aspectos emocionais, comportamentais "
        "ou relacionais, exigindo suporte terapêutico para favorecer o desenvolvimento "
        "de habilidades adaptativas e funcionais. A avaliação clínica sugere que, "
        "embora não haja indicativos de psicopatologia, a intervenção é pertinente "
        "para promoção do bem-estar, prevenção de dificuldades futuras e apoio ao "
        "desenvolvimento global.",
    )

    add_section_title(doc, "Evolução")
    add_section_text(
        doc,
        "Desde o início do acompanhamento, o(a) paciente tem demonstrado avanços "
        "compatíveis com os objetivos terapêuticos estabelecidos. Observa-se "
        "progressiva ampliação da capacidade de expressão emocional, melhor "
        "compreensão de situações internas e externas e maior tolerância a "
        "frustrações e contrariedades. Há indícios de fortalecimento do vínculo "
        "terapêutico, o que tem favorecido maior abertura para o diálogo, "
        "elaboração de vivências e desenvolvimento de estratégias de enfrentamento. "
        "Em casos infantis, o uso de recursos lúdicos, histórias sociais e "
        "brincadeiras tem promovido maior engajamento e expressão simbólica. "
        "Para adolescentes e adultos, observa-se maior clareza na identificação "
        "de sentimentos e pensamento reflexivo sobre padrões de comportamento, "
        "relações interpessoais e tomada de decisões.",
    )

    add_section_title(doc, "Programação Terapêutica Atual")
    especialidades_encontradas = [esp.upper() for esp in patient_data["especialidades"]]

    if any("PSICOTERAPIA" in esp.upper() for esp in patient_data["especialidades"]):
        add_section_text(
            doc,
            "A psicoterapia segue com o objetivo de promover o autoconhecimento, "
            "fortalecer a autoestima e desenvolver recursos internos para lidar com "
            "desafios emocionais e comportamentais. São utilizadas estratégias adequadas "
            "à faixa etária, tais como: escuta ativa, ludoterapia, mediação simbólica, "
            "reestruturação cognitiva, treino de habilidades sociais e técnicas de "
            "regulação emocional. A abordagem terapêutica está centrada nas necessidades "
            "atuais do(a) paciente, com foco na construção de estratégias saudáveis "
            "para resolução de conflitos internos, desenvolvimento da autonomia emocional "
            "e aprimoramento das relações interpessoais. O processo psicoterapêutico é "
            "conduzido respeitando o ritmo individual, com observação contínua da evolução "
            "e ajustes nas intervenções conforme a resposta do(a) paciente.",
        )

    if any("ABA" in esp or "TERAPIA ABA" in esp for esp in especialidades_encontradas):
        add_section_title(doc, "Terapia ABA", Pt(16), Pt(8))
        add_section_text(
            doc,
            "A programação atual visa fortalecer comportamentos funcionais, "
            "ampliar a comunicação, promover autonomia nas rotinas e reduzir "
            "comportamentos de oposição, fuga ou auto estimulação. São utilizados "
            "programas personalizados de ensino por tentativas discretas, ensino "
            "naturalístico e treino de habilidades sociais.",
        )

    if any("TERAPIA OCUPACIONAL" in esp or "OCUPACIONAL" in esp for esp in especialidades_encontradas):
        add_section_title(doc, "Terapia Ocupacional", Pt(16), Pt(8))
        add_section_text(
            doc,
            "As intervenções atuais priorizam o desenvolvimento da independência "
            "em atividades de vida diária (AVDs), o planejamento motor e a integração "
            "sensorial. São propostas atividades lúdicas e funcionais, com adaptações "
            "conforme a faixa etária, para favorecer o desempenho ocupacional global.",
        )

    if any("FONOAUDIOLOGIA" in esp or "FONO" in esp for esp in especialidades_encontradas):
        add_section_title(doc, "Fonoaudiologia", Pt(16), Pt(8))
        add_section_text(
            doc,
            "O foco terapêutico envolve o aperfeiçoamento da linguagem oral e/ou "
            "alternativa, melhora na compreensão e expressão verbal, bem como o "
            "desenvolvimento das habilidades fonológicas e comunicativas. A intervenção "
            "considera o nível atual de linguagem e o contexto escolar, familiar e social da paciente.",
        )

    if any("PSICOMOTRICIDADE" in esp or "PSICOMOTOR" in esp for esp in especialidades_encontradas):
        add_section_title(doc, "Psicomotricidade", Pt(16), Pt(8))
        add_section_text(
            doc,
            "O trabalho psicomotor tem como meta promover o domínio do corpo no espaço, "
            "controle postural, lateralidade e coordenação em diferentes níveis. As sessões "
            "envolvem jogos, circuitos e desafios motores com objetivos específicos para "
            "aprimorar a integração sensório-motora.",
        )

    if any("PSICOPEDAGOGIA" in esp or "PEDAGOG" in esp for esp in especialidades_encontradas):
        add_section_title(doc, "Psicopedagogia", Pt(16), Pt(8))
        add_section_text(
            doc,
            "A atuação psicopedagógica busca estimular habilidades cognitivas e acadêmicas, "
            "com estratégias personalizadas para desenvolver leitura, escrita, lógica "
            "matemática e resolução de problemas. O plano também inclui o fortalecimento "
            "da autoestima escolar e o apoio no planejamento e organização do tempo.",
        )

    add_section_title(doc, "Considerações Finais")
    add_section_text(
        doc,
        "Recomenda-se a continuidade do processo terapêutico, com participação "
        "ativa da família e alinhamento com a rede de apoio para promover a "
        "generalização dos avanços obtidos em consultório. A psicoterapia tem se "
        "mostrado um espaço importante de escuta, acolhimento e construção de "
        "recursos para a promoção da saúde mental.",
    )

    add_fixed_signature_section(doc, "Brasília, data da assinatura digital.")

    # add_signature_section(doc, "tipico")
    build_signature_block_grid(doc, patient_data["especialidades"], cols=2)

    filename = f"Relatório_Típico_{patient_data['info']['nome'].replace(' ', '_')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
