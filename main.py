import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import os
import sys
from collections import defaultdict

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

class MedicalReportGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("Gerador de Relatórios FUSEX")
        self.root.geometry("900x700")
        
        # Configurar ícone se existir
        try:
            self.root.iconbitmap(resource_path("icone.ico"))
        except:
            pass  # Ignora se não tiver ícone
        
        # Variáveis
        self.excel_file = None
        self.report_type = tk.StringVar(value="PNE")
        self.data = None
        
        self.setup_ui()
        
    def setup_ui(self):
        # Frame principal
        main_frame = ttk.Frame(self.root, padding="15")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Título com estilo
        title_label = ttk.Label(main_frame, text="🏥 Gerador de Relatórios FUSEX", 
                               font=("Arial", 18, "bold"))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 25))
        
        # Seleção do tipo de relatório
        type_frame = ttk.LabelFrame(main_frame, text="📋 Tipo de Relatório", padding="15")
        type_frame.grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 15))
        
        ttk.Radiobutton(type_frame, text="🧩 PNE (Portador de Necessidades Especiais)", 
                       variable=self.report_type, value="PNE").pack(anchor=tk.W, pady=5)
        ttk.Radiobutton(type_frame, text="👤 Típico", 
                       variable=self.report_type, value="TIPICO").pack(anchor=tk.W, pady=5)
        
        # Seleção de arquivo
        file_frame = ttk.LabelFrame(main_frame, text="📁 Arquivo Excel", padding="15")
        file_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 15))
        
        # Frame interno para organizar melhor
        file_inner_frame = ttk.Frame(file_frame)
        file_inner_frame.pack(fill=tk.X)
        
        self.file_label = ttk.Label(file_inner_frame, text="Nenhum arquivo selecionado", 
                                   foreground="gray")
        self.file_label.pack(side=tk.LEFT, padx=(0, 15))
        
        # Botões de arquivo
        btn_frame = ttk.Frame(file_inner_frame)
        btn_frame.pack(side=tk.RIGHT)
        
        ttk.Button(btn_frame, text="📁 Selecionar Arquivo", 
                  command=self.select_file).pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(btn_frame, text="📄 Usar Exemplo", 
                  command=self.load_example).pack(side=tk.LEFT)
        
        # Preview dos dados
        preview_frame = ttk.LabelFrame(main_frame, text="👁️ Preview dos Dados", padding="15")
        preview_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 15))
        
        # Treeview para mostrar os dados
        columns = ("Nome", "Data Nascimento", "Responsável", "Especialidade", "Mês Referência")
        self.tree = ttk.Treeview(preview_frame, columns=columns, show="headings", height=12)
        
        for col in columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=150)
        
        scrollbar_v = ttk.Scrollbar(preview_frame, orient=tk.VERTICAL, command=self.tree.yview)
        scrollbar_h = ttk.Scrollbar(preview_frame, orient=tk.HORIZONTAL, command=self.tree.xview)
        self.tree.configure(yscrollcommand=scrollbar_v.set, xscrollcommand=scrollbar_h.set)
        
        self.tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        scrollbar_v.grid(row=0, column=1, sticky=(tk.N, tk.S))
        scrollbar_h.grid(row=1, column=0, sticky=(tk.W, tk.E))
        
        preview_frame.grid_rowconfigure(0, weight=1)
        preview_frame.grid_columnconfigure(0, weight=1)
        
        # Botões de ação
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, columnspan=3, pady=(15, 0))
        
        ttk.Button(button_frame, text="🚀 Gerar Relatórios", 
                  command=self.generate_reports).pack(side=tk.LEFT, padx=(0, 15))
        ttk.Button(button_frame, text="🗑️ Limpar", 
                  command=self.clear_data).pack(side=tk.LEFT, padx=(0, 15))
        ttk.Button(button_frame, text="❓ Ajuda", 
                  command=self.show_help).pack(side=tk.LEFT)
        
        # Status bar
        self.status_var = tk.StringVar(value="✅ Pronto para uso")
        status_bar = ttk.Label(main_frame, textvariable=self.status_var, 
                              relief=tk.SUNKEN, anchor=tk.W, padding="5")
        status_bar.grid(row=5, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(15, 0))
        
        # Configurar redimensionamento
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(2, weight=1)
        main_frame.rowconfigure(3, weight=1)
    
    def load_example(self):
        """Carrega o arquivo de exemplo fusex_tipico.xlsx"""
        try:
            example_file = resource_path("fusex_tipico.xlsx")
            if os.path.exists(example_file):
                self.excel_file = example_file
                self.data = pd.read_excel(example_file)
                
                # Verificar colunas
                required_columns = ['NOME', 'DATA DE NASCIMENTO', 'RESPONSÁVEL', 
                                  'ESPECIALIDADE', 'MÊS DE REFERÊNCIA']
                
                missing_columns = [col for col in required_columns if col not in self.data.columns]
                if missing_columns:
                    messagebox.showerror("Erro", f"Colunas obrigatórias não encontradas no exemplo: {', '.join(missing_columns)}")
                    return
                
                self.file_label.config(text="📄 fusex_tipico.xlsx (exemplo)", foreground="blue")
                self.load_preview()
                self.status_var.set(f"📄 Exemplo carregado: {len(self.data)} registros")
            else:
                messagebox.showerror("Erro", "Arquivo de exemplo não encontrado!")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao carregar exemplo: {str(e)}")
    
    def show_help(self):
        """Mostra janela de ajuda"""
        help_window = tk.Toplevel(self.root)
        help_window.title("Ajuda - Gerador de Relatórios FUSEX")
        help_window.geometry("600x400")
        help_window.transient(self.root)
        help_window.grab_set()
        
        help_text = """
        📋 COMO USAR O GERADOR DE RELATÓRIOS FUSEX
        
        1️⃣ PREPARAR PLANILHA EXCEL:
           • Colunas obrigatórias:
             - NOME
             - DATA DE NASCIMENTO
             - RESPONSÁVEL
             - ESPECIALIDADE
             - MÊS DE REFERÊNCIA
        
        2️⃣ SELECIONAR TIPO DE RELATÓRIO:
           • PNE: Para portadores de necessidades especiais
           • Típico: Para desenvolvimento típico
        
        3️⃣ CARREGAR DADOS:
           • Use "Selecionar Arquivo" para sua planilha
           • Ou "Usar Exemplo" para testar
        
        4️⃣ GERAR RELATÓRIOS:
           • Clique em "Gerar Relatórios"
           • Escolha pasta de destino
           • Aguarde o processamento
        
        ⚙️ ESPECIALIDADES SUPORTADAS:
           • Terapia ABA
           • Psicoterapia
           • Terapia Ocupacional
           • Fonoaudiologia
           • Psicomotricidade
           • Psicopedagogia
        
        📧 SUPORTE: Qualquer dúvida, entre em contato!
        """
        
        text_widget = tk.Text(help_window, wrap=tk.WORD, padx=20, pady=20)
        text_widget.insert(tk.END, help_text)
        text_widget.config(state=tk.DISABLED)
        text_widget.pack(fill=tk.BOTH, expand=True)
        
        ttk.Button(help_window, text="✅ Entendi", 
                  command=help_window.destroy).pack(pady=10)
    
    def add_table_borders(self, table):
        """Adiciona bordas à tabela de cabeçalho"""
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    tcBorders.append(border)
                
                tcPr.append(tcBorders)
    
    def create_header_table(self, doc, patient_data, convenio="FUSEX"):
        """Cria a tabela de cabeçalho padronizada"""
        table = doc.add_table(rows=6, cols=1)
        self.add_table_borders(table)
        
        especialidades_unicas = list(set(patient_data['especialidades']))
        especialidades_str = ", ".join(especialidades_unicas)
        
        header_data = [
            f"Nome: {patient_data['info']['nome']}",
            f"Data de Nascimento: {patient_data['info']['data_nascimento']}",
            f"Responsável: {patient_data['info']['responsavel']}",
            f"Convênio: {convenio}",
            f"Especialidade: {especialidades_str}",
            f"Mês de referência: {patient_data['info']['mes_referencia']}"
        ]
        
        for i, data in enumerate(header_data):
            cell = table.cell(i, 0)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if ':' in data:
                parts = data.split(':', 1)
                title_run = paragraph.add_run(parts[0] + ': ')
                title_run.bold = True
                paragraph.add_run(parts[1])
            else:
                run = paragraph.add_run(data)
                run.bold = True
        
        # Parágrafo em branco simples para separar da próxima seção
        doc.add_paragraph()
        return table
    
    def add_section_title(self, doc, title, space_before=Pt(24), space_after=Pt(12)):
        """Adiciona título de seção padronizado"""
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.space_before = space_before
        fmt.space_after = space_after
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(title)
        run.bold = True
        return p
    
    def add_section_text(self, doc, text, space_before=Pt(12), space_after=Pt(12)):
        """Adiciona texto de seção padronizado"""
        p = doc.add_paragraph(text)
        fmt = p.paragraph_format
        fmt.space_before = space_before
        fmt.space_after = space_after
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p
    
    def add_specialty_section(self, doc, specialty_name, evolution_text, programming_text=None, is_evolution=True):
        """Adiciona seção de especialidade (para relatórios PNE)"""
        # Título da especialidade
        if is_evolution:
            self.add_section_title(doc, specialty_name, Pt(18), Pt(8))
        else:
            self.add_section_title(doc, specialty_name, Pt(16), Pt(8))
        
        # Texto de evolução
        self.add_section_text(doc, evolution_text)
        
        # Se há texto de programação, adicionar também
        if programming_text:
            self.add_section_text(doc, programming_text)
    
    def add_signature_section(self, doc, signature_type="tipico"):
        """Adiciona seção de assinatura padronizada"""
        # Parágrafo em branco para espaçamento antes da assinatura
        doc.add_paragraph()
        
        # Linha de assinatura
        p = doc.add_paragraph("_" * 50)
        fmt = p.paragraph_format
        fmt.space_before = Pt(24)
        fmt.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if signature_type == "pne":
            # Para relatórios PNE
            resp_p = doc.add_paragraph("Responsável técnica(o)")
            resp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Para relatórios típicos
            psico_p = doc.add_paragraph("Responsável técnica(o)")
            psico_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return p
        
    def select_file(self):
        file_path = filedialog.askopenfilename(
            title="Selecionar arquivo Excel",
            filetypes=[("Excel files", "*.xlsx *.xls")]
        )
        
        if file_path:
            try:
                self.excel_file = file_path
                self.data = pd.read_excel(file_path)
                
                # Verificar se tem as colunas necessárias
                required_columns = ['NOME', 'DATA DE NASCIMENTO', 'RESPONSÁVEL', 
                                  'ESPECIALIDADE', 'MÊS DE REFERÊNCIA']
                
                missing_columns = [col for col in required_columns if col not in self.data.columns]
                if missing_columns:
                    messagebox.showerror("Erro", f"Colunas obrigatórias não encontradas: {', '.join(missing_columns)}")
                    return
                
                self.file_label.config(text=f"📁 {os.path.basename(file_path)}", foreground="green")
                self.load_preview()
                self.status_var.set(f"📊 Arquivo carregado: {len(self.data)} registros")
                
            except Exception as e:
                messagebox.showerror("Erro", f"Erro ao carregar arquivo: {str(e)}")
    
    def load_preview(self):
        # Limpar dados anteriores
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        # Carregar novos dados
        for _, row in self.data.iterrows():
            # Formatar data de nascimento
            data_nasc = row['DATA DE NASCIMENTO']
            try:
                if pd.notna(data_nasc):
                    if hasattr(data_nasc, 'strftime'):
                        data_nasc = data_nasc.strftime('%d/%m/%Y')
                    else:
                        data_nasc = str(data_nasc)
                        if '/' in data_nasc:
                            data_nasc = data_nasc.split(' ')[0]
                else:
                    data_nasc = "Data não informada"
            except:
                data_nasc = str(data_nasc) if pd.notna(data_nasc) else "Data não informada"
            
            self.tree.insert("", tk.END, values=(
                str(row['NOME']) if pd.notna(row['NOME']) else "Nome não informado",
                data_nasc,
                str(row['RESPONSÁVEL']) if pd.notna(row['RESPONSÁVEL']) else "Não informado",
                str(row['ESPECIALIDADE']) if pd.notna(row['ESPECIALIDADE']) else "Não informado",
                str(row['MÊS DE REFERÊNCIA']) if pd.notna(row['MÊS DE REFERÊNCIA']) else "Não informado"
            ))
    
    def clear_data(self):
        self.excel_file = None
        self.data = None
        self.file_label.config(text="Nenhum arquivo selecionado", foreground="gray")
        for item in self.tree.get_children():
            self.tree.delete(item)
        self.status_var.set("✅ Pronto para uso")
    
    def generate_reports(self):
        if self.data is None:
            messagebox.showerror("Erro", "Nenhum arquivo carregado! Use 'Selecionar Arquivo' ou 'Usar Exemplo'.")
            return
        
        # Agrupar dados por paciente
        patient_data = defaultdict(lambda: {
            'info': {},
            'especialidades': []
        })
        
        for _, row in self.data.iterrows():
            nome = row['NOME']
            if not patient_data[nome]['info']:
                # Formatar data de nascimento
                data_nasc = row['DATA DE NASCIMENTO']
                try:
                    if pd.notna(data_nasc):
                        if hasattr(data_nasc, 'strftime'):
                            data_nasc = data_nasc.strftime('%d/%m/%Y')
                        else:
                            data_nasc = str(data_nasc)
                            if '/' in data_nasc:
                                data_nasc = data_nasc.split(' ')[0]
                    else:
                        data_nasc = "Data não informada"
                except:
                    data_nasc = str(data_nasc) if pd.notna(data_nasc) else "Data não informada"
                
                patient_data[nome]['info'] = {
                    'nome': str(row['NOME']),
                    'data_nascimento': data_nasc,
                    'responsavel': str(row['RESPONSÁVEL']) if pd.notna(row['RESPONSÁVEL']) else "Não informado",
                    'mes_referencia': str(row['MÊS DE REFERÊNCIA']) if pd.notna(row['MÊS DE REFERÊNCIA']) else "Não informado"
                }
            
            # Adicionar especialidade se não for nula
            if pd.notna(row['ESPECIALIDADE']):
                patient_data[nome]['especialidades'].append(str(row['ESPECIALIDADE']))
        
        # Selecionar diretório de saída
        output_dir = filedialog.askdirectory(title="Selecionar pasta para salvar relatórios")
        if not output_dir:
            return
        
        try:
            self.status_var.set("🔄 Gerando relatórios...")
            self.root.update()
            
            report_count = 0
            for nome, data in patient_data.items():
                if self.report_type.get() == "PNE":
                    self.generate_pne_report(data, output_dir)
                else:
                    self.generate_tipico_report(data, output_dir)
                report_count += 1
                
                # Atualizar progresso
                self.status_var.set(f"🔄 Gerando... {report_count}/{len(patient_data)}")
                self.root.update()
            
            self.status_var.set(f"✅ {report_count} relatórios gerados com sucesso!")
            messagebox.showinfo("Sucesso! 🎉", 
                              f"✅ {report_count} relatórios gerados com sucesso!\n\n"
                              f"📁 Pasta: {output_dir}")
            
        except Exception as e:
            self.status_var.set("❌ Erro ao gerar relatórios")
            messagebox.showerror("Erro", f"Erro ao gerar relatórios: {str(e)}")
    
    def generate_pne_report(self, patient_data, output_dir):
        """Gera relatório PNE com espaçamento otimizado"""
        # Criar documento
        try:
            papel_timbrado = resource_path("papel timbrado.docx")
            doc = Document(papel_timbrado)
        except:
            doc = Document()
            # Cabeçalho básico como papel timbrado
            header_p = doc.add_paragraph()
            header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header_p.add_run("CLÍNICA MÉDICA - PNE")
            run.bold = True
            run.font.size = Pt(16)
            # Parágrafo em branco após cabeçalho
            doc.add_paragraph()
        
        # Cabeçalho
        self.create_header_table(doc, patient_data, "Fusex PNE")
        
        # Hipótese Diagnóstica
        self.add_section_title(doc, "Hipótese Diagnóstica")
        self.add_section_text(doc, 
            "Transtorno do Espectro Autista, conforme critérios do DSM-5. "
            "Apresentando prejuízos significativos na comunicação social e "
            "comportamentos restritos e repetitivos, necessitando de apoio "
            "substancial em múltiplas áreas do desenvolvimento.")
        
        # Verificar especialidades
        especialidades_encontradas = [esp.upper() for esp in patient_data['especialidades']]
        
        # Evolução por especialidade
        if any("ABA" in esp or "TERAPIA ABA" in esp for esp in especialidades_encontradas):
            self.add_specialty_section(doc, "Terapia ABA",
                "A paciente tem apresentado evolução gradual no manejo de "
                "comportamentos desafiadores e na aquisição de habilidades adaptativas. "
                "Destaca-se o progresso na capacidade de seguir instruções, maior "
                "participação em atividades estruturadas e aumento da comunicação "
                "funcional, com boa aceitação aos programas propostos e resposta "
                "positiva ao reforço positivo.", is_evolution=True)
        
        if any("PSICOTERAPIA" in esp for esp in especialidades_encontradas):
            self.add_specialty_section(doc, "Psicoterapia",
                "Observa-se maior abertura da paciente ao vínculo terapêutico, "
                "com avanços na expressão emocional, identificação de sentimentos "
                "e melhora na tolerância a frustrações. Há desenvolvimento de "
                "estratégias internas de enfrentamento e maior consciência sobre "
                "suas próprias emoções e comportamentos, adequadas à sua faixa etária.", is_evolution=True)
        
        if any("TERAPIA OCUPACIONAL" in esp or "OCUPACIONAL" in esp for esp in especialidades_encontradas):
            self.add_specialty_section(doc, "Terapia Ocupacional",
                "Houve progresso no desempenho ocupacional, especialmente nas "
                "áreas de autorregulação, coordenação motora e autonomia nas "
                "atividades diárias. A paciente apresenta melhor organização "
                "sensorial e maior engajamento em tarefas funcionais, tanto em "
                "contextos lúdicos quanto nas rotinas cotidianas.", is_evolution=True)
        
        if any("FONOAUDIOLOGIA" in esp or "FONO" in esp for esp in especialidades_encontradas):
            self.add_specialty_section(doc, "Fonoaudiologia",
                "Verifica-se avanço significativo nas habilidades comunicativas, "
                "seja por meio da fala, linguagem alternativa ou recursos expressivos "
                "e receptivos. A paciente demonstra melhor compreensão de comandos, "
                "maior intenção comunicativa e expansão do vocabulário funcional, "
                "além de avanços na articulação e fluência, conforme a necessidade individual.", is_evolution=True)
        
        if any("PSICOMOTRICIDADE" in esp or "PSICOMOTOR" in esp for esp in especialidades_encontradas):
            self.add_specialty_section(doc, "Psicomotricidade",
                "A evolução psicomotora inclui melhora na coordenação global e fina, "
                "organização espacial e equilíbrio. A paciente demonstra maior "
                "consciência corporal e controle motor, com progressos que refletem "
                "positivamente no comportamento, na atenção e na interação social "
                "durante as atividades terapêuticas.", is_evolution=True)
        
        if any("PSICOPEDAGOGIA" in esp or "PEDAGOG" in esp for esp in especialidades_encontradas):
            self.add_specialty_section(doc, "Psicopedagogia",
                "A paciente apresentou melhora na atenção, concentração e interesse "
                "por atividades que envolvem linguagem, raciocínio lógico e habilidades "
                "acadêmicas. Observa-se avanço na memória de trabalho, organização do "
                "pensamento e capacidade de seguir sequências, contribuindo para o "
                "desempenho escolar ou acadêmico, conforme a faixa etária.", is_evolution=True)
        
        # Programação Terapêutica Atual
        self.add_section_title(doc, "Programação Terapêutica Atual", Pt(30), Pt(12))
        
        # Programação por especialidade
        if any("ABA" in esp or "TERAPIA ABA" in esp for esp in especialidades_encontradas):
            self.add_section_title(doc, "Terapia ABA", Pt(16), Pt(8))
            self.add_section_text(doc,
                "A programação atual visa fortalecer comportamentos funcionais, "
                "ampliar a comunicação, promover autonomia nas rotinas e reduzir "
                "comportamentos de oposição, fuga ou auto estimulação. São utilizados "
                "programas personalizados de ensino por tentativas discretas, ensino "
                "naturalístico e treino de habilidades sociais.")
        
        if any("PSICOTERAPIA" in esp for esp in especialidades_encontradas):
            self.add_section_title(doc, "Psicoterapia", Pt(16), Pt(8))
            self.add_section_text(doc,
                "Os objetivos terapêuticos incluem promover o autoconhecimento, "
                "a regulação emocional, o desenvolvimento da autoestima e o enfrentamento "
                "saudável de desafios, utilizando abordagens adequadas à idade "
                "(brincadeiras simbólicas, recursos visuais, técnicas cognitivas, entre outras).")
        
        if any("TERAPIA OCUPACIONAL" in esp or "OCUPACIONAL" in esp for esp in especialidades_encontradas):
            self.add_section_title(doc, "Terapia Ocupacional", Pt(16), Pt(8))
            self.add_section_text(doc,
                "As intervenções atuais priorizam o desenvolvimento da independência "
                "em atividades de vida diária (AVDs), o planejamento motor e a integração "
                "sensorial. São propostas atividades lúdicas e funcionais, com adaptações "
                "conforme a faixa etária, para favorecer o desempenho ocupacional global.")
        
        if any("FONOAUDIOLOGIA" in esp or "FONO" in esp for esp in especialidades_encontradas):
            self.add_section_title(doc, "Fonoaudiologia", Pt(16), Pt(8))
            self.add_section_text(doc,
                "O foco terapêutico envolve o aperfeiçoamento da linguagem oral e/ou "
                "alternativa, melhora na compreensão e expressão verbal, bem como o "
                "desenvolvimento das habilidades fonológicas e comunicativas. A intervenção "
                "considera o nível atual de linguagem e o contexto escolar, familiar e social da paciente.")
        
        if any("PSICOMOTRICIDADE" in esp or "PSICOMOTOR" in esp for esp in especialidades_encontradas):
            self.add_section_title(doc, "Psicomotricidade", Pt(16), Pt(8))
            self.add_section_text(doc,
                "O trabalho psicomotor tem como meta promover o domínio do corpo no espaço, "
                "controle postural, lateralidade e coordenação em diferentes níveis. As sessões "
                "envolvem jogos, circuitos e desafios motores com objetivos específicos para "
                "aprimorar a integração sensório-motora.")
        
        if any("PSICOPEDAGOGIA" in esp or "PEDAGOG" in esp for esp in especialidades_encontradas):
            self.add_section_title(doc, "Psicopedagogia", Pt(16), Pt(8))
            self.add_section_text(doc,
                "A atuação psicopedagógica busca estimular habilidades cognitivas e acadêmicas, "
                "com estratégias personalizadas para desenvolver leitura, escrita, lógica "
                "matemática e resolução de problemas. O plano também inclui o fortalecimento "
                "da autoestima escolar e o apoio no planejamento e organização do tempo.")
        
        # Considerações Finais
        self.add_section_title(doc, "Considerações Finais", Pt(30), Pt(12))
        self.add_section_text(doc,
            "A paciente segue em acompanhamento com evolução positiva. O trabalho "
            "conjunto entre as especialidades tem favorecido ganhos significativos e "
            "generalização das habilidades desenvolvidas para diferentes ambientes. "
            "Recomendamos a continuidade do atendimento terapêutico e o envolvimento "
            "da família e/ou escola no processo.")
        
        self.add_section_text(doc,
            "Nos colocamos à disposição para esclarecimentos sobre o processo "
            "terapêutico, bem como para oferecer orientações e suporte sempre que "
            "necessário, respeitando os limites éticos da atuação clínica.")
        
        # Assinatura
        self.add_signature_section(doc, "pne")
        
        # Salvar documento
        filename = f"Relatório_PNE_{patient_data['info']['nome'].replace(' ', '_')}.docx"
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
    
    def generate_tipico_report(self, patient_data, output_dir):
        """Gera relatório Típico com espaçamento otimizado"""
        # Criar documento
        try:
            papel_timbrado = resource_path("papel timbrado.docx")
            doc = Document(papel_timbrado)
        except:
            doc = Document()
            # Cabeçalho básico como papel timbrado
            header_p = doc.add_paragraph()
            header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header_p.add_run("CLÍNICA MÉDICA")
            run.bold = True
            run.font.size = Pt(16)
            # Parágrafo em branco após cabeçalho
            doc.add_paragraph()
        
        # Cabeçalho
        self.create_header_table(doc, patient_data, "FUSEX")
        
        # Hipótese Diagnóstica
        self.add_section_title(doc, "Hipótese Diagnóstica")
        self.add_section_text(doc,
            "O paciente apresenta características compatíveis com o desenvolvimento "
            "típico, sem indicação, até o momento, de transtornos diagnosticáveis "
            "conforme os manuais classificatórios vigentes (CID-11/DSM-5). As demandas "
            "observadas referem-se a dificuldades específicas no enfrentamento de "
            "situações do cotidiano, que podem envolver aspectos emocionais, comportamentais "
            "ou relacionais, exigindo suporte terapêutico para favorecer o desenvolvimento "
            "de habilidades adaptativas e funcionais. A avaliação clínica sugere que, "
            "embora não haja indicativos de psicopatologia, a intervenção é pertinente "
            "para promoção do bem-estar, prevenção de dificuldades futuras e apoio ao "
            "desenvolvimento global.")
        
        # Evolução
        self.add_section_title(doc, "Evolução")
        self.add_section_text(doc,
            "Desde o início do acompanhamento, o(a) paciente tem demonstrado avanços "
            "compatíveis com os objetivos terapêuticos estabelecidos. Observa-se "
            "progressiva ampliação da capacidade de expressão emocional, melhor "
            "compreensão de situações internas e externas e maior tolerância a "
            "frustrações e contrariedades. Há indícios de fortalecimento do vínculo "
            "terapêutico, o que tem favorecido maior abertura para o diálogo, "
            "elaboração de vivências e desenvolvimento de estratégias de enfrentamento. "
            "Em casos infantis, o uso de recursos lúdicos, histórias sociais e "
            "brincadeiras tem promovido maior engajamento e expressão simbólica. "
            "Para adolescentes e adultos, observa-se maior clareza na identificação "
            "de sentimentos e pensamento reflexivo sobre padrões de comportamento, "
            "relações interpessoais e tomada de decisões.")
        
        # Programação Terapêutica Atual
        self.add_section_title(doc, "Programação Terapêutica Atual")
        
        # Verificar se tem psicoterapia nas especialidades
        especialidades_encontradas = [esp.upper() for esp in patient_data['especialidades']]
        
        if any("PSICOTERAPIA" in esp.upper() for esp in patient_data['especialidades']):
            self.add_section_text(doc,
                "A psicoterapia segue com o objetivo de promover o autoconhecimento, "
                "fortalecer a autoestima e desenvolver recursos internos para lidar com "
                "desafios emocionais e comportamentais. São utilizadas estratégias adequadas "
                "à faixa etária, tais como: escuta ativa, ludoterapia, mediação simbólica, "
                "reestruturação cognitiva, treino de habilidades sociais e técnicas de "
                "regulação emocional. A abordagem terapêutica está centrada nas necessidades "
                "atuais do(a) paciente, com foco na construção de estratégias saudáveis "
                "para resolução de conflitos internos, desenvolvimento da autonomia emocional "
                "e aprimoramento das relações interpessoais. O processo psicoterapêutico é "
                "conduzido respeitando o ritmo individual, com observação contínua da evolução "
                "e ajustes nas intervenções conforme a resposta do(a) paciente.")
        
        if any("ABA" in esp or "TERAPIA ABA" in esp for esp in especialidades_encontradas):
            self.add_section_title(doc, "Terapia ABA", Pt(16), Pt(8))
            self.add_section_text(doc,
                "A programação atual visa fortalecer comportamentos funcionais, "
                "ampliar a comunicação, promover autonomia nas rotinas e reduzir "
                "comportamentos de oposição, fuga ou auto estimulação. São utilizados "
                "programas personalizados de ensino por tentativas discretas, ensino "
                "naturalístico e treino de habilidades sociais.")
        
        if any("TERAPIA OCUPACIONAL" in esp or "OCUPACIONAL" in esp for esp in especialidades_encontradas):
            self.add_section_title(doc, "Terapia Ocupacional", Pt(16), Pt(8))
            self.add_section_text(doc,
                "As intervenções atuais priorizam o desenvolvimento da independência "
                "em atividades de vida diária (AVDs), o planejamento motor e a integração "
                "sensorial. São propostas atividades lúdicas e funcionais, com adaptações "
                "conforme a faixa etária, para favorecer o desempenho ocupacional global.")
        
        if any("FONOAUDIOLOGIA" in esp or "FONO" in esp for esp in especialidades_encontradas):
            self.add_section_title(doc, "Fonoaudiologia", Pt(16), Pt(8))
            self.add_section_text(doc,
                "O foco terapêutico envolve o aperfeiçoamento da linguagem oral e/ou "
                "alternativa, melhora na compreensão e expressão verbal, bem como o "
                "desenvolvimento das habilidades fonológicas e comunicativas. A intervenção "
                "considera o nível atual de linguagem e o contexto escolar, familiar e social da paciente.")
        
        if any("PSICOMOTRICIDADE" in esp or "PSICOMOTOR" in esp for esp in especialidades_encontradas):
            self.add_section_title(doc, "Psicomotricidade", Pt(16), Pt(8))
            self.add_section_text(doc,
                "O trabalho psicomotor tem como meta promover o domínio do corpo no espaço, "
                "controle postural, lateralidade e coordenação em diferentes níveis. As sessões "
                "envolvem jogos, circuitos e desafios motores com objetivos específicos para "
                "aprimorar a integração sensório-motora.")
        
        if any("PSICOPEDAGOGIA" in esp or "PEDAGOG" in esp for esp in especialidades_encontradas):
            self.add_section_title(doc, "Psicopedagogia", Pt(16), Pt(8))
            self.add_section_text(doc,
                "A atuação psicopedagógica busca estimular habilidades cognitivas e acadêmicas, "
                "com estratégias personalizadas para desenvolver leitura, escrita, lógica "
                "matemática e resolução de problemas. O plano também inclui o fortalecimento "
                "da autoestima escolar e o apoio no planejamento e organização do tempo.")
        
        # Considerações Finais
        self.add_section_title(doc, "Considerações Finais")
        self.add_section_text(doc,
            "Recomenda-se a continuidade do processo terapêutico, com participação "
            "ativa da família e alinhamento com a rede de apoio para promover a "
            "generalização dos avanços obtidos em consultório. A psicoterapia tem se "
            "mostrado um espaço importante de escuta, acolhimento e construção de "
            "recursos para a promoção da saúde mental.")
        
        # Assinatura
        self.add_signature_section(doc, "tipico")
        
        # Salvar documento
        filename = f"Relatório_Típico_{patient_data['info']['nome'].replace(' ', '_')}.docx"
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)

def main():
    root = tk.Tk()
    app = MedicalReportGenerator(root)
    root.mainloop()

if __name__ == "__main__":
    main()